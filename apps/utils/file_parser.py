"""
文件解析工具模块
支持解析多种格式的需求文档，提取文本内容
"""

import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# 支持的文件类型
SUPPORTED_FILE_TYPES = {
    '.docx': 'Microsoft Word文档',
    '.md': 'Markdown文档',
    '.txt': '纯文本文件',
    '.pdf': 'PDF文档',
}

# 允许的文件扩展名
ALLOWED_EXTENSIONS = list(SUPPORTED_FILE_TYPES.keys())

# 文件大小限制（50MB）
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

def is_allowed_file(filename):
    """
    检查文件是否为允许的类型
    
    Args:
        filename: 文件名
        
    Returns:
        bool: 是否允许上传
    """
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in [ext[1:] for ext in ALLOWED_EXTENSIONS]

def is_file_size_allowed(file_size):
    """
    检查文件大小是否在允许范围内
    
    Args:
        file_size: 文件大小（字节）
        
    Returns:
        bool: 是否允许上传
    """
    return file_size <= MAX_FILE_SIZE

def get_human_readable_size(bytes_size):
    """
    将字节转换为人类可读的文件大小
    
    Args:
        bytes_size: 字节数
        
    Returns:
        str: 人类可读的文件大小
    """
    if bytes_size < 1024:
        return f"{bytes_size} B"
    elif bytes_size < 1024 * 1024:
        return f"{bytes_size / 1024:.2f} KB"
    else:
        return f"{bytes_size / (1024 * 1024):.2f} MB"

def get_file_extension(filename):
    """
    获取文件扩展名（小写）
    
    Args:
        filename: 文件名
        
    Returns:
        str: 扩展名（包含点号）
    """
    if '.' in filename:
        return '.' + filename.rsplit('.', 1)[1].lower()
    return ''

def parse_docx(file_path):
    """
    解析Word文档(.docx)
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本内容
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        content = []
        
        # 提取段落
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                content.append(text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    content.append(' | '.join(row_text))
        
        return '\n\n'.join(content)
        
    except ImportError:
        logger.error("python-docx库未安装，请安装: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"解析DOCX文件失败: {e}")
        raise

def parse_markdown(file_path):
    """
    解析Markdown文档(.md)
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本内容
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content
        
    except UnicodeDecodeError:
        # 尝试其他编码
        with open(file_path, 'r', encoding='gbk') as f:
            content = f.read()
        return content
    except Exception as e:
        logger.error(f"解析Markdown文件失败: {e}")
        raise

def parse_txt(file_path):
    """
    解析纯文本文件(.txt)
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本内容
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content
        
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='gbk') as f:
            content = f.read()
        return content
    except Exception as e:
        logger.error(f"解析TXT文件失败: {e}")
        raise

def parse_pdf(file_path):
    """
    解析PDF文档(.pdf)
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本内容
    """
    try:
        # 尝试使用PyPDF2解析
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            content = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content.append(text.strip())
            
            return '\n\n'.join(content)
            
        except ImportError:
            logger.warning("PyPDF2库未安装，尝试使用pdfplumber")
            
            # 尝试使用pdfplumber解析
            try:
                import pdfplumber
                
                with pdfplumber.open(file_path) as pdf:
                    content = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            content.append(text.strip())
                
                return '\n\n'.join(content)
                
            except ImportError:
                logger.error("PyPDF2和pdfplumber库均未安装，请安装其中一个: pip install PyPDF2 或 pip install pdfplumber")
                raise ImportError("需要安装PyPDF2或pdfplumber库来解析PDF文件")
                
    except Exception as e:
        logger.error(f"解析PDF文件失败: {e}")
        raise

def parse_file(file_path):
    """
    根据文件类型解析文件内容
    
    Args:
        file_path: 文件路径
        
    Returns:
        tuple: (content, file_type)
            content: 提取的文本内容
            file_type: 文件类型描述
    """
    ext = get_file_extension(file_path)
    
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型: {ext}。支持的类型: {', '.join(SUPPORTED_FILE_TYPES.values())}")
    
    logger.info(f"开始解析文件: {file_path}, 类型: {SUPPORTED_FILE_TYPES.get(ext, ext)}")
    
    if ext == '.docx':
        content = parse_docx(file_path)
    elif ext == '.md':
        content = parse_markdown(file_path)
    elif ext == '.txt':
        content = parse_txt(file_path)
    elif ext == '.pdf':
        content = parse_pdf(file_path)
    else:
        raise ValueError(f"未实现的文件类型解析: {ext}")
    
    logger.info(f"文件解析完成，内容长度: {len(content)} 字符")
    return content, SUPPORTED_FILE_TYPES.get(ext, ext)

def extract_text_from_uploaded_file(uploaded_file):
    """
    从上传的文件对象中提取文本内容
    
    Args:
        uploaded_file: Django上传的文件对象
        
    Returns:
        tuple: (content, file_name, file_type)
            content: 提取的文本内容
            file_name: 原始文件名
            file_type: 文件类型描述
    """
    filename = uploaded_file.name
    
    # 验证文件类型
    if not is_allowed_file(filename):
        ext = get_file_extension(filename)
        raise ValueError(f"不支持的文件类型: {ext}。支持的类型: {', '.join(SUPPORTED_FILE_TYPES.values())}")
    
    # 验证文件大小
    if hasattr(uploaded_file, 'size'):
        file_size = uploaded_file.size
        if not is_file_size_allowed(file_size):
            raise ValueError(f"文件大小超出限制。当前大小: {get_human_readable_size(file_size)}，最大允许: {get_human_readable_size(MAX_FILE_SIZE)}")
    
    # 创建临时文件
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix=get_file_extension(filename)) as temp_file:
        for chunk in uploaded_file.chunks():
            temp_file.write(chunk)
        temp_path = temp_file.name
    
    try:
        content, file_type = parse_file(temp_path)
        return content, filename, file_type
    finally:
        # 清理临时文件
        if os.path.exists(temp_path):
            os.unlink(temp_path)

def extract_text_from_multiple_files(files):
    """
    从多个上传文件中提取文本内容
    
    Args:
        files: 文件对象列表
        
    Returns:
        tuple: (combined_content, file_info_list)
            combined_content: 所有文件内容合并后的文本
            file_info_list: 文件信息列表，包含filename和file_type
    """
    combined_content = []
    file_info_list = []
    
    for uploaded_file in files:
        content, filename, file_type = extract_text_from_uploaded_file(uploaded_file)
        combined_content.append(f"=== {filename} ===\n\n{content}")
        file_info_list.append({
            'filename': filename,
            'file_type': file_type,
            'content_length': len(content)
        })
    
    return '\n\n'.join(combined_content), file_info_list
